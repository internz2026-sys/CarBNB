"""Generates the CarBNB Tester Handoff Guide as a .docx file.

One-shot artifact — not part of the app build. Run with:
    python scripts/generate-tester-guide.py
Output: CarBNB-Tester-Handoff-Guide.docx at the project root.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches


# ──────────────────────────── Styling helpers ────────────────────────────


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x00, 0x3D, 0x9B)
    if subtitle:
        sub = doc.add_paragraph()
        sub_run = sub.add_run(subtitle)
        sub_run.italic = True
        sub_run.font.size = Pt(10)
        sub_run.font.color.rgb = RGBColor(0x60, 0x60, 0x70)
    doc.add_paragraph()


def h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x00, 0x3D, 0x9B)


def h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x1F, 0x3A)


def h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1A, 0x1F, 0x3A)


def body(doc, segments, bullet=False, number=False, size=11):
    if bullet:
        p = doc.add_paragraph(style="List Bullet")
    elif number:
        p = doc.add_paragraph(style="List Number")
    else:
        p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    if isinstance(segments, str):
        segments = [(segments, False)]
    for text, bold in segments:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)


def callout(doc, label, text):
    """Inline emphasis callout — e.g., 'Expected: ...' / 'Tip: ...'."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.3
    lab = p.add_run(f"{label} ")
    lab.bold = True
    lab.font.size = Pt(10.5)
    lab.font.color.rgb = RGBColor(0x00, 0x3D, 0x9B)
    txt = p.add_run(text)
    txt.font.size = Pt(10.5)
    txt.italic = True


def step(doc, text, expected=None):
    """Numbered step with an optional expected-result line."""
    body(doc, text, number=True)
    if expected:
        callout(doc, "Expected:", expected)


def checkbox(doc, text):
    body(doc, f"[ ]  {text}", bullet=False)


def hr(doc):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C7CDDB")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ──────────────────────────── Content ────────────────────────────


def build():
    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    add_title(
        doc,
        "carBNB — Tester Handoff Guide",
        "A hands-on walkthrough of every role and flow in the MVP. Work through the sections in order — each builds on the previous one.",
    )
    hr(doc)

    # ──────────────── 0. Before you start ────────────────
    h1(doc, "0. Before you start")

    h3(doc, "0.1 What you'll need")
    body(doc, "A modern browser (Chrome, Firefox, Edge, Safari) — ideally with a private/incognito window so you can hold multiple sessions side-by-side.", bullet=True)
    body(doc, "The admin credentials (ask the project owner — typically admin@carbnb.com + the shared password).", bullet=True)
    body(doc, "A scratch email domain you can make up on the spot — Supabase Auth has 'Confirm email' turned OFF in this environment, so any unique fake address works (e.g., tester-host-01@example.com, tester-customer-01@example.com).", bullet=True)
    body(doc, "Two test photos (any JPG/PNG/WebP under 5 MB each) and one test PDF (under 5 MB) for the OR/CR upload — a screenshot works fine.", bullet=True)

    h3(doc, "0.2 Where to test")
    body(
        doc,
        [
            ("The live production URL is ", False),
            ("https://car-bnb-eta.vercel.app", True),
            (". This is what you will demo. All the steps in this guide run there unless noted.", False),
        ],
    )

    h3(doc, "0.3 Mental model — three roles")
    body(
        doc,
        [
            ("Admin", True),
            (": the platform operator. Owns verifications, bookings mid-lifecycle, cash payments, and settings. Lives under URLs like /dashboard, /owners, /bookings, /settings.", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("Host", True),
            (": someone who owns a car and earns from renting it out. Signs up via the Host tab. Lives under /host/*.", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("Customer (Renter)", True),
            (": someone who books a car. Signs up via the Customer tab. Lives under /account and books from the public /listings pages.", False),
        ],
        bullet=True,
    )

    h3(doc, "0.4 How to keep three sessions open at once")
    body(doc, "You will flip between roles a lot. Use: (a) your default browser logged in as admin, (b) a private/incognito window as the host, (c) a second browser (or a second private window) as the customer. This saves logging in and out all day.", size=11)

    hr(doc)

    # ──────────────── 1. Admin walkthrough ────────────────
    h1(doc, "1. Admin walkthrough")
    body(doc, "Start here. The admin touches every corner of the platform, so this is the fastest way to learn the data model.")

    h3(doc, "1.1 Log in as admin")
    step(
        doc,
        [
            ("Go to ", False),
            ("https://car-bnb-eta.vercel.app/login", True),
            (". Use either tab — admin accounts bypass tab validation.", False),
        ],
        expected="After login you land on /dashboard.",
    )

    h3(doc, "1.2 The dashboard")
    step(
        doc,
        "Look at the four tiles across the top: Total Owners, Active Listings, Total Revenue, Pending Approvals. These are live numbers — they come straight from the database.",
        expected="Numbers roughly match the seeded dataset. The Total Revenue card shows a Platform Fee (15%) breakdown and an Owner Payouts total.",
    )
    step(
        doc,
        "Scroll down. On the left is the Verification Queue (up to 3 pending owners + 3 pending listings). On the right is the Active & Upcoming Bookings table.",
    )
    step(
        doc,
        "Click any card in the Verification Queue.",
        expected="It deep-links to the owner or listing detail page where you'd actually approve it. Don't approve anything yet — just navigate back.",
    )
    step(
        doc,
        "Click any row in the Active & Upcoming Bookings table (the vehicle cell).",
        expected="Lands on /bookings/[id] with the full booking detail. Navigate back to the dashboard.",
    )

    h3(doc, "1.3 /owners — host directory")
    step(
        doc,
        "Click 'Owners' in the sidebar.",
        expected="You see all hosts across all statuses. The filter chips at the top are All / Pending / Verified / Suspended.",
    )
    step(
        doc,
        "Click into any VERIFIED owner.",
        expected="You see their full profile, uploaded documents (ID + driver's license), and an 'Edit' button. You can view the uploaded documents via the 1-hour signed-URL link.",
    )
    step(
        doc,
        "Go back and click Add Owner (top-right). Fill in a fake name + email + phone + address. Save.",
        expected="New owner appears in the directory with status PENDING. Nothing is visible to the public yet.",
    )
    step(
        doc,
        "Open the owner you just created. Click the Approve button.",
        expected="Status flips to VERIFIED. An activity-log entry was silently written under the hood.",
    )

    h3(doc, "1.4 /car-listings — vehicle directory")
    step(
        doc,
        "Click 'Car Listings' in the sidebar.",
        expected="See all vehicles across all statuses. Status chips: Pending / Active / Suspended / Booked.",
    )
    step(
        doc,
        "Click Add Listing. Select the owner you just approved in 1.3. Fill in plate number (invent one like 'TST-001'), brand, model, year, color, location (e.g., 'Iloilo City'), transmission (Automatic), fuel type, seating, daily rate (e.g., 2500). Save.",
        expected="You land on /car-listings/[id]/edit with status PENDING_APPROVAL.",
    )
    step(
        doc,
        "On that edit page: upload at least 2 photos (scroll to the Photos section, click 'Choose photo'). Upload the OR/CR document (test PDF works). In Weekly Availability, toggle Mon–Fri on with start 08:00 / end 18:00; toggle Sat/Sun off. Click Save Schedule.",
    )
    step(
        doc,
        "In Date-Specific Exceptions, add a blackout date for next Wednesday with reason 'Maintenance'. Click Add.",
        expected="The exception appears in the list below with a red icon.",
    )
    step(
        doc,
        "Scroll back up and click Approve.",
        expected="Status flips to ACTIVE. The listing is now publicly browsable.",
    )

    h3(doc, "1.5 /availability — fleet-wide overview")
    step(
        doc,
        "Click 'Availability' in the sidebar.",
        expected="Read-only summary of every listing's weekly schedule. No actions here — this is a reference view.",
    )

    h3(doc, "1.6 /calendar — fleet calendar")
    step(
        doc,
        "Click 'Calendar' in the sidebar.",
        expected="A month-view calendar with colored dates. Blue = booked ranges. Red = blocked exceptions.",
    )
    step(
        doc,
        "Change the vehicle dropdown from 'All Vehicles' to the listing you just created in 1.4. The calendar should show your Wednesday blackout in red.",
    )
    step(
        doc,
        "Click the blocked Wednesday. On the right panel, you should see the Blocked card with the reason 'Maintenance' and the vehicle name.",
    )

    h3(doc, "1.7 /bookings — booking management")
    step(
        doc,
        "Click 'Bookings' in the sidebar.",
        expected="Full booking list with filters (All / Pending / Confirmed / Ongoing / Completed / Cancelled / Rejected / Unpaid). Summary tiles up top.",
    )
    step(
        doc,
        "Click New Booking.",
        expected="A form for creating an admin-initiated booking.",
    )
    step(
        doc,
        "Pick a customer (any seeded one) and the listing you approved in 1.4. Pick a date range that doesn't cross your Wednesday blackout. Confirm the fee preview shows. Add a note. Submit.",
        expected="New booking appears as CONFIRMED immediately (admin-created bookings skip PENDING). You land on its detail page.",
    )
    step(
        doc,
        "On the booking detail page, click Start Rental.",
        expected="Status becomes ONGOING. 'Rental Started' timestamp appears.",
    )
    step(
        doc,
        "Click Mark as Paid.",
        expected="Dialog asks for optional notes. Submit. Payment status flips to PAID with a green banner.",
    )
    step(
        doc,
        "Click Complete Rental.",
        expected="Status becomes COMPLETED. 'Rental Completed' timestamp appears. The booking now counts in the host's Total Earnings on their dashboard.",
    )

    h3(doc, "1.8 /customers — renter directory")
    step(
        doc,
        "Click 'Customers' in the sidebar.",
        expected="List of every customer with their booking count and last booking date.",
    )
    step(
        doc,
        "Click View on any customer with bookings.",
        expected="Detail page with contact info, booking-count breakdown (completed / cancelled), lifetime spend tile, and full booking history.",
    )

    h3(doc, "1.9 /settings — platform settings")
    step(
        doc,
        "Click 'Settings' in the sidebar.",
        expected="Form with: Commission Rate (%), Security Deposit (₱), Auto-Approve toggle, Require Owner Confirmation toggle, Minimum Booking Notice (hrs).",
    )
    step(
        doc,
        "Change commission from 15 to 20. Click Save Settings.",
        expected="Green 'Settings saved.' banner. Footer shows 'Last saved by <your admin email> · <timestamp>'.",
    )
    step(
        doc,
        "Go back to /dashboard. The 'Platform Fee (X%)' label on the revenue card now reads 20%.",
        expected="The amounts don't change — historical bookings keep their locked-in rate. Only the label reflects the new setting.",
    )
    step(
        doc,
        "Return to /settings and change commission back to 15. Save.",
    )

    h3(doc, "1.10 Logout")
    step(
        doc,
        "Click your avatar top-right → Log out.",
        expected="You're back on the landing page, not logged in.",
    )

    hr(doc)

    # ──────────────── 2. Host walkthrough ────────────────
    h1(doc, "2. Host walkthrough")
    body(doc, "Now play the role of a new host signing up. Use a fresh browser window/session so you don't disturb the admin session.")

    h3(doc, "2.1 Sign up as a host")
    step(
        doc,
        "Go to /signup. Click the Host tab. Use a new email like tester-host-01@example.com, any full name (e.g., 'Test Host'), and a password of 8+ characters. Submit.",
        expected="Redirect to /login?signedUp=host.",
    )
    step(
        doc,
        "Log in with the credentials you just used, on the Host tab.",
        expected="You land on /host/dashboard. It shows a 'Your host account is under review' locked screen with an amber shield icon. None of the three tiles render.",
    )
    step(
        doc,
        "Try clicking 'My Cars' or 'My Bookings' in the nav.",
        expected="Both are hidden from the nav for pending hosts — the only visible item is Dashboard.",
    )
    step(
        doc,
        "Try typing /host/cars directly into the URL.",
        expected="Bounces you back to /host/dashboard (page-level redirect for non-verified hosts).",
    )

    h3(doc, "2.2 Admin approves the host")
    step(
        doc,
        "In your admin tab, go to /owners, filter to Pending, and find the host you just signed up.",
    )
    step(
        doc,
        "Click into their profile → Approve.",
        expected="Status becomes VERIFIED.",
    )

    h3(doc, "2.3 Host is now verified")
    step(
        doc,
        "Back in the host tab, refresh /host/dashboard.",
        expected="The three tiles now render: Active Listings (0), Upcoming Bookings (0), Total Earnings (₱0). Nav now shows Dashboard / My Cars / My Bookings.",
    )

    h3(doc, "2.4 Host creates a listing")
    step(
        doc,
        "Click 'My Cars' → 'Add Listing'. Invent a plate (e.g., 'HOST-201'), fill in vehicle details, daily rate around 3000. Submit.",
        expected="You land on /host/cars/[id]/edit. The top chip says PENDING_APPROVAL and the subtitle reads 'Waiting on admin approval'.",
    )
    step(
        doc,
        "Upload at least one photo. Upload OR/CR. Toggle weekly availability for a few days. Add one blackout exception. Save each section as you go.",
    )

    h3(doc, "2.5 Admin approves the listing")
    step(
        doc,
        "In admin tab: /car-listings → filter Pending → click into the new listing → Approve.",
    )
    step(
        doc,
        "Host tab: refresh /host/cars/[id]/edit.",
        expected="Chip now says ACTIVE. Subtitle changes to 'Changes save instantly and do not require re-approval.'",
    )

    h3(doc, "2.6 Host edits price (free post-approval)")
    step(
        doc,
        "On the host edit page, change Daily Rate from 3000 to 3200. Save.",
        expected="Success. Status stays ACTIVE — no re-approval bounce.",
    )

    hr(doc)

    # ──────────────── 3. Customer walkthrough ────────────────
    h1(doc, "3. Customer walkthrough")
    body(doc, "Third persona. Again use a fresh session.")

    h3(doc, "3.1 Sign up as a customer")
    step(
        doc,
        "Go to /signup. Click the Customer tab. Use tester-customer-01@example.com, any name, password.",
        expected="Redirect to /login?signedUp=customer.",
    )
    step(
        doc,
        "Log in on the Customer tab.",
        expected="Lands on /account. You see 'No active bookings' empty state.",
    )

    h3(doc, "3.2 Browse public listings")
    step(
        doc,
        "Click 'Browse cars' in the user menu top-right, or go to /listings.",
        expected="Public listings page. You should see the host's new car from Section 2 among the results.",
    )
    step(
        doc,
        "Try the search box (e.g., search for the brand you created). Try the location filter.",
    )
    step(
        doc,
        "Click into the host's listing you created.",
        expected="Detail page with photos, host name with Verified badge, a sticky 'Reserve Now' bar at the bottom.",
    )

    h3(doc, "3.3 Book a car")
    step(
        doc,
        "Click Reserve Now. Dialog opens with a range calendar. Dates that are blocked (the host's exception) should be greyed out.",
    )
    step(
        doc,
        "Pick a pickup + return date 2 days apart, avoiding the blackout date.",
        expected="A live fee preview appears: N days × daily rate, platform fee, owner payout, customer total.",
    )
    step(
        doc,
        "Deliberately pick a range that crosses the blackout date.",
        expected="A red warning 'Your selection crosses unavailable days' appears. The Reserve button is disabled until you fix the range.",
    )
    step(
        doc,
        "Fix the range, click Reserve.",
        expected="Booking is created. You land back on /account with a new PENDING booking card showing a BK-XXXXXX reference, dates, and car info.",
    )

    h3(doc, "3.4 Host sees and accepts the booking")
    step(
        doc,
        "In the host tab, go to /host/bookings.",
        expected="The new PENDING booking appears.",
    )
    step(
        doc,
        "Click into it. See the 'Respond to this request' card with Accept / Reject buttons.",
    )
    step(
        doc,
        "Click Accept Booking.",
        expected="Status flips to CONFIRMED. The Accept/Reject card disappears. Admin will own the rest of the lifecycle.",
    )

    h3(doc, "3.5 Customer self-cancel (only works on PENDING)")
    step(
        doc,
        "Back in the customer tab, create a second booking on the same listing (different dates).",
    )
    step(
        doc,
        "On /account, click the new PENDING booking. You should see a 'Cancel booking' button.",
    )
    step(
        doc,
        "Click Cancel. Confirm.",
        expected="Status flips to CANCELLED. Customer-initiated cancellations are allowed only while the booking is still PENDING.",
    )

    hr(doc)

    # ──────────────── 4. End-to-end happy path ────────────────
    h1(doc, "4. End-to-end happy path (the demo story)")
    body(
        doc,
        "If you only have 10 minutes to show this platform, run this sequence. It touches every major feature and tells a clean narrative.",
    )

    body(
        doc,
        [
            ("Act 1 — Host signs up. ", True),
            ("Show /signup → Host tab → submit. Show the locked 'under review' screen. Key point: platform doesn't trust-by-default.", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("Act 2 — Admin verifies. ", True),
            ("Switch to admin. Show /owners → Pending queue → click approve. Mention: 'This is where the admin would check uploaded ID and license docs — the signed URLs are 1-hour, never public.'", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("Act 3 — Host lists a car. ", True),
            ("Switch to host. Show /host/dashboard now unlocked. Create a listing, upload a photo, set availability, submit. Show the PENDING_APPROVAL chip.", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("Act 4 — Admin approves the car. ", True),
            ("Admin tab → /car-listings → click approve. Mention: 'No car appears in the public marketplace until it passes review.'", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("Act 5 — Customer discovers and books. ", True),
            ("Customer tab → /listings → find the car → Reserve Now. Show the availability-aware calendar + live fee preview. Submit.", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("Act 6 — Host accepts. ", True),
            ("Host tab → /host/bookings → Accept. Walk through the reason dropdown on Reject to show the audit-trail mechanism.", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("Act 7 — Admin runs the trip. ", True),
            ("Admin tab → /bookings/[id] → Start Rental → Mark as Paid (cash) → Complete Rental. Explain: 'Cash is first-class on this platform — every handoff is a recorded event with admin, timestamp, method, notes.'", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("Act 8 — Everyone sees the outcome. ", True),
            ("Host sees Total Earnings tile updated. Admin sees dashboard revenue tiles updated. Customer sees the booking under Past Bookings. Close by showing /settings and mentioning 'commission is configurable, new rates only apply to new bookings, historical totals are locked.'", False),
        ],
        bullet=True,
    )

    hr(doc)

    # ──────────────── 5. Gotchas & edge cases ────────────────
    h1(doc, "5. Gotchas & edge cases worth knowing")

    body(
        doc,
        [
            ("Tab mismatch. ", True),
            ("If a host tries to log in via the Customer tab (or vice versa), the platform rejects the attempt with a clear error and signs them out. Admin accounts are exempt and can use either tab. Try this deliberately in prod so you know what the error looks like.", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("Login email is preserved on error. ", True),
            ("Typing a valid email but a wrong password does NOT blank the email field on retry. Small UX thing but users notice.", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("Pending hosts can't create. ", True),
            ("Even if a pending host manually navigates to /host/cars/new, the server will refuse to create a listing. UI blocks this too, but the server is the real gate.", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("Suspended hosts. ", True),
            ("If an admin suspends a host, the host sees a red-shield 'Your host account is suspended' locked view instead of their dashboard. Test this to see what it looks like, then re-verify so your demo data stays clean.", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("Commission lock-in. ", True),
            ("Changing the commission in /settings doesn't retroactively change stored bookings. Each booking keeps its original rate. Dashboard's 'Platform Fee (X%)' label is the live rate; per-booking detail shows the computed historical rate from stored values.", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("Cross-host isolation. ", True),
            ("Host A trying to open /host/cars/[host-b-id]/edit gets a 404. Similarly for bookings. Safe to assume all host data is scoped per-owner.", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("Inclusive-day billing. ", True),
            ("A booking from May 4 to May 6 counts as 3 days, not 2. This matches what the calendar visually highlights and the local P2P rental convention.", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("Cancellation reasons are structured. ", True),
            ("When an admin or host cancels/rejects, they pick from a preset dropdown (customer no-show, documents not verified, vehicle unavailable, duplicate booking, other). Note is required for 'other'. All of this is written to the audit log.", False),
        ],
        bullet=True,
    )

    hr(doc)

    # ──────────────── 6. Known limitations ────────────────
    h1(doc, "6. Known limitations — what NOT to demo")
    body(
        doc,
        "These areas are planned for later tiers. If asked, acknowledge they're deferred, not broken.",
    )

    body(
        doc,
        [
            ("/accounting page", True),
            (" still reads mock data. Don't open it in a demo.", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("/reports page", True),
            (" is a placeholder. Skip it.", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("No email notifications yet.", True),
            (" Booking confirmations, approval notices, etc. are all in-app only for now. Supabase 'Confirm email' is also OFF so fake signups don't require inbox access.", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("Payment is cash-only.", True),
            (" No PayMongo / GCash / card integration yet — admin manually records 'Paid' with optional notes. This matches the current PH informal-rental reality and is intentional for now.", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("No messaging between host and renter.", True),
            (" Coordination happens out-of-band (phone/SMS) right now.", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("No reviews or ratings yet.", True),
            (" The stars you see on the landing page featured cards are decorative.", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("Admin dashboard visuals", True),
            (" could use more polish — numbers are right, layout is functional but not final.", False),
        ],
        bullet=True,
    )

    hr(doc)

    # ──────────────── 7. Pre-demo checklist ────────────────
    h1(doc, "7. Pre-demo checklist (day-of)")

    checkbox(doc, "All three tabs/windows open: admin, host, customer.")
    checkbox(doc, "Commission rate in /settings is set to 15% (default).")
    checkbox(doc, "At least one listing is ACTIVE with good photos and clear availability.")
    checkbox(doc, "No leftover 'test' PENDING bookings cluttering the admin dashboard — clean them up before going live.")
    checkbox(doc, "You know the admin password and have it within reach.")
    checkbox(doc, "You have a throwaway email ready to sign up a new host or customer live during the demo (makes the story feel real).")
    checkbox(doc, "You've read Section 5 (Gotchas) so nothing surprises you.")
    checkbox(doc, "You've practiced Section 4 (happy path) at least twice end-to-end.")

    hr(doc)

    # ──────────────── 8. If something goes wrong ────────────────
    h1(doc, "8. If something goes wrong mid-demo")

    body(
        doc,
        [
            ("A page 500s. ", True),
            ("Reload once. If it repeats, check the Vercel deployment status or ping the project owner. Don't try to 'fix it' live.", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("A booking won't confirm. ", True),
            ("Usually because the date range overlaps an existing booking or blocked day. Check the listing's availability on /calendar first.", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("Login fails. ", True),
            ("The most common cause is being on the wrong tab (customer email on Host tab). The error message will tell you — just switch tabs and retry.", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("Uploaded photo doesn't appear. ", True),
            ("Refresh the page. Supabase Storage occasionally takes a second to propagate.", False),
        ],
        bullet=True,
    )
    body(
        doc,
        [
            ("Session unexpectedly logs out. ", True),
            ("The Supabase session cookies are 1-hour. If you left a tab open for a while, expect to re-login.", False),
        ],
        bullet=True,
    )

    hr(doc)
    body(
        doc,
        [
            ("End of guide. Questions, findings, or unexpected behaviors → write them down and send to the project owner. Happy testing.", False),
        ],
    )

    out_path = Path(__file__).resolve().parent.parent / "CarBNB-Tester-Handoff-Guide.docx"
    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    build()
