"""Generates the CarBNB Product & Operational Digest as a .docx file.

One-shot artifact generator — not part of the app build. Run with:
    python scripts/generate-digest.py
Output: CarBNB-Product-Digest.docx at the project root.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches


def add_title(doc, text, subtitle=None):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(text)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x00, 0x3D, 0x9B)
    if subtitle:
        sub = doc.add_paragraph()
        sub_run = sub.add_run(subtitle)
        sub_run.italic = True
        sub_run.font.size = Pt(10)
        sub_run.font.color.rgb = RGBColor(0x60, 0x60, 0x70)
    doc.add_paragraph()


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x00, 0x3D, 0x9B)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1A, 0x1F, 0x3A)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1A, 0x1F, 0x3A)


def add_body(doc, segments):
    """Add a paragraph. segments is a list of (text, bold_bool) tuples, OR a plain string."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.35
    if isinstance(segments, str):
        segments = [(segments, False)]
    for text, bold in segments:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(11)


def add_bullet(doc, segments):
    """Add a bullet-point. segments same shape as add_body."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    if isinstance(segments, str):
        segments = [(segments, False)]
    for text, bold in segments:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(11)


def add_hr(doc):
    # Thin horizontal rule via a bottom border on an empty paragraph
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C7CDDB")
    pBdr.append(bottom)
    pPr.append(pBdr)


def build():
    doc = Document()

    # Tighten default body margins a touch
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_title(
        doc,
        "carBNB — Product & Operational Digest",
        "Prepared from the live site (car-bnb-eta.vercel.app) and the product source. Written for non-technical readers.",
    )
    add_hr(doc)

    # ---- 1. The Value Proposition ----
    add_heading(doc, "1. The Value Proposition", level=1)
    add_body(
        doc,
        [
            ("carBNB is not a rental listings board and not a fleet-rental company. It positions itself as ", False),
            ("“The Curated Engine”", True),
            (" — a peer-to-peer marketplace where every host, every vehicle, and every trip passes through a hospitality-grade review before it becomes publicly visible. The product language is deliberate: ", False),
            ("“editorial-quality profile,” “curated income,” “high-trust by design,”", True),
            (" and ", False),
            ("“Sharing built with hospitality, not friction.”", True),
            (" The visible brand promise is three words: ", False),
            ("Verified hosts · Transparent payouts · Support-led trips.", True),
        ],
    )
    add_body(
        doc,
        [
            ("The niche, in plain language: ", False),
            ("an Airbnb-style take on local car-sharing.", True),
            (" Where most Philippine P2P rental activity lives in Facebook groups and Viber threads, carBNB offers a curated storefront — featured cars, verified owners, structured booking references, and a transparent fee split. The site avoids rent-a-car vocabulary entirely (“dealerships,” “fleet,” “commodity inventory”) and instead borrows from hospitality and editorial publishing (“trending now,” “featured host drive,” “community signal”). The positioning is clear: vehicles are ", False),
            ("presented", True),
            (", not ", False),
            ("listed", True),
            ("; hosts are ", False),
            ("curated", True),
            (", not ", False),
            ("onboarded", True),
            (".", False),
        ],
    )
    add_body(
        doc,
        [
            ("This premium framing supports a concrete commercial differentiation: hosts get a ", False),
            ("calmer, more professional", True),
            (" way to monetize idle vehicles; renters get ", False),
            ("confident discovery", True),
            (" instead of ad-hoc negotiation.", False),
        ],
    )

    add_hr(doc)

    # ---- 2. Operational Mechanics ----
    add_heading(doc, "2. Operational Mechanics — How the Curated Engine Functions", level=1)
    add_body(
        doc,
        [
            ("The platform is built around a ", False),
            ("double-gated intake", True),
            (". Nothing appears to the public until it passes explicit admin review.", False),
        ],
    )

    add_heading(doc, "From “idle asset” to “listed vehicle”:", level=2)
    add_bullet(
        doc,
        [
            ("A prospective host signs up via the Host tab. Their account is created in ", False),
            ("PENDING", True),
            (" state and they immediately land on a locked “Your host account is under review” dashboard — they can log in, but cannot list, edit, or transact.", False),
        ],
    )
    add_bullet(
        doc,
        [
            ("The host (or an admin on their behalf) uploads the three required documents: ", False),
            ("government ID", True),
            (", ", False),
            ("driver’s license", True),
            (", and ", False),
            ("OR/CR (vehicle registration papers)", True),
            (". All three land in private cloud storage and are only visible to admins via short-lived (1-hour) signed URLs.", False),
        ],
    )
    add_bullet(
        doc,
        [
            ("An admin reviews the application in the Verification Queue on the main dashboard and promotes the owner to ", False),
            ("VERIFIED", True),
            (". Only now does the host unlock the full workspace.", False),
        ],
    )
    add_bullet(
        doc,
        [
            ("The host submits a car through a simplified form (no owner field — identity is implicit). The listing is created in ", False),
            ("PENDING_APPROVAL", True),
            (". The host then uploads ", False),
            ("photos (up to 8)", True),
            (", the ", False),
            ("vehicle’s OR/CR", True),
            (", and sets ", False),
            ("weekly availability", True),
            (" + ", False),
            ("blackout dates", True),
            (".", False),
        ],
    )
    add_bullet(
        doc,
        [
            ("An admin reviews the listing, verifies the documents, and promotes it to ", False),
            ("ACTIVE", True),
            (". From this moment — and not before — the vehicle appears in the public browse page and landing-page featured rail.", False),
        ],
    )
    add_bullet(
        doc,
        [
            ("Post-approval edits (price tweaks, description, photo swaps) take effect ", False),
            ("immediately, without re-review", True),
            (" — matching the Turo/Airbnb convention that gives hosts pricing agility while preserving the initial vetting.", False),
        ],
    )

    add_heading(doc, "Payout transparency (directly from the code):", level=2)
    add_body(
        doc,
        [
            ("A pure helper function, ", False),
            ("calculateBookingAmount(dailyPrice, pickupDate, returnDate, commissionRate)", True),
            (", is the single source of truth for every quote and every stored booking total. It uses ", False),
            ("inclusive calendar-day counting", True),
            (" — May 4 to May 6 counts as three days — matching what the date picker visually highlights. The helper returns a three-way split: ", False),
            ("totalAmount → platformFee → ownerPayout", True),
            (".", False),
        ],
    )
    add_body(
        doc,
        [
            ("The platform commission (default ", False),
            ("15%", True),
            (") lives in a single settings row editable by admins from /settings. Crucially, the rate a booking was created at is ", False),
            ("permanently stored on that booking", True),
            (". Raising the commission tomorrow does not retroactively change yesterday’s earnings — the host’s payout on a completed trip is locked the moment the customer hits “Reserve.” Both the customer’s dialog (before submit) and the host’s booking detail screen show the identical breakdown: ", False),
            ("Customer total → Platform fee (X%) → Owner payout.", True),
            (" Same numbers, both sides, all the way through.", False),
        ],
    )

    add_hr(doc)

    # ---- 3. User Experience ----
    add_heading(doc, "3. User Experience — Two Personas", level=1)

    add_heading(doc, "The Verified Host", level=2)
    add_body(
        doc,
        [
            ("The host’s workspace is built around the single metaphor of ", False),
            ("managing an asset", True),
            (", not fulfilling tasks. They land on /host/dashboard with three tiles: ", False),
            ("Active Listings", True),
            (", ", False),
            ("Upcoming Bookings", True),
            (", and ", False),
            ("Total Earnings", True),
            (" (a running sum across completed trips). Each is clickable and deep-links into the relevant section.", False),
        ],
    )
    add_body(
        doc,
        [
            ("From ", False),
            ("My Cars", True),
            (", the host sees their fleet as a photo-first grid with status chips (Active, Pending, Suspended). They can add a new listing, swap photos, replace the OR/CR document, adjust the weekly rentable window (e.g., Monday–Friday 08:00–18:00), and add one-off exceptions — a blocked weekend for personal use, or a forced-available holiday. Price changes take effect instantly; there is no waiting period.", False),
        ],
    )
    add_body(
        doc,
        [
            ("From ", False),
            ("My Bookings", True),
            (", the host sees a filtered queue of requests on ", False),
            ("their", True),
            (" cars. For a PENDING request they get two actions — ", False),
            ("Accept Booking", True),
            (" (flips to CONFIRMED) or ", False),
            ("Reject", True),
            (" (dialog with a preset reason dropdown plus an optional free-text note, required if “Other”). Once confirmed, the request leaves the host’s action bar — the platform operator owns the rest of the lifecycle (trip start, trip completion, cash collection), matching the Turo model. This division keeps hosts focused on request acceptance while guaranteeing one authoritative source of truth for mid-trip events.", False),
        ],
    )
    add_body(
        doc,
        [
            ("Throughout, every host action is ", False),
            ("ownership-scoped at the data layer", True),
            (". A host cannot — under any URL manipulation — view, edit, or respond to another host’s listing or booking. A denied request returns a 404 or a “you cannot modify” error; there is no leakage.", False),
        ],
    )

    add_heading(doc, "The Explorer Renter", level=2)
    add_body(
        doc,
        [
            ("The renter’s journey is intentionally ", False),
            ("short, visual, and trust-forward", True),
            (". The homepage opens with a featured host drive, a “Fully Verified” badge, and live community-signal numbers (verified hosts, trips in motion, metro coverage) pulled straight from the database. The three-step renter narrative — ", False),
            ("Locate → Book → Drive", True),
            (" — is reinforced across the landing, the browse page, and the detail view.", False),
        ],
    )
    add_body(
        doc,
        [
            ("On /listings, the renter searches by keyword or location. On a car’s detail page they see the photo gallery, transmission/fuel/seat specs, the host’s name + verified status, and a ", False),
            ("90-day availability map", True),
            (" where dates that are booked, blocked by the weekly schedule, or blacked-out by an exception are all greyed out from the picker in one pass.", False),
        ],
    )
    add_body(
        doc,
        [
            ("When they click ", False),
            ("Reserve Now", True),
            (", a dialog opens with a range calendar and a ", False),
            ("live fee preview", True),
            (": the daily-rate × days math, plus the platform fee and owner payout as informational context. If the selected range accidentally crosses a blocked day — for example, picking a date after a maintenance window — a red warning appears ", False),
            ("before they can submit", True),
            (" (“Your selection crosses unavailable days”), gently nudging them to split into two bookings. This client-side guardrail is the same check the server enforces on submit, so a rejected booking is extraordinarily rare.", False),
        ],
    )
    add_body(
        doc,
        [
            ("On submit, a structured reference number is generated (", False),
            ("BK-A3X9K2", True),
            (" format — readable, non-sequential so total volume isn’t leaked) and the booking enters PENDING. The customer sees it instantly on /account under ", False),
            ("Upcoming & Active", True),
            (", alongside a status chip. They can self-cancel while it’s still PENDING — once a host or admin confirms, further changes require operator involvement.", False),
        ],
    )

    add_hr(doc)

    # ---- 4. Trust & Safety ----
    add_heading(doc, "4. Trust & Safety Infrastructure", level=1)
    add_body(
        doc,
        [
            ("Trust is enforced in ", False),
            ("four stacked layers", True),
            (", each independently verifiable.", False),
        ],
    )
    add_body(
        doc,
        [
            ("Layer 1 — Identity at the middleware. ", True),
            ("Every page request is intercepted by the root middleware (proxy.ts). For any admin, customer, or host URL, the middleware performs two checks: (a) a valid Supabase session exists, and (b) a corresponding row exists in the domain database for that email. The ", False),
            ("database is authoritative", True),
            ("; auth metadata alone is never trusted, because auth metadata can be written directly by a client using the public key. This single decision eliminates the most common identity-spoofing vector in Next.js apps.", False),
        ],
    )
    add_body(
        doc,
        [
            ("Layer 2 — Role and status at the server action. ", True),
            ("Every write to the database goes through a gate: requireAdmin(), requireHost() (verified-only; pending and suspended hosts are rejected), or ownership-scope checks (listing.ownerId === current.host.id). Pending hosts can log in and see their locked dashboard — but the server action layer will refuse any mutation they attempt, even if a client tries to bypass the UI.", False),
        ],
    )
    add_body(
        doc,
        [
            ("Layer 3 — Document verification workflow. ", True),
            ("Sensitive documents (government ID, driver’s license, vehicle OR/CR) never touch the public internet. They live in private cloud storage buckets and are surfaced to admins only through ", False),
            ("1-hour signed URLs", True),
            (" generated per-view. Photos — which are promotional, not identifying — are the only public asset. A host without an approved OR/CR on file cannot get their listing approved; the admin’s review page puts the document viewer directly next to the Approve button.", False),
        ],
    )
    add_body(
        doc,
        [
            ("Layer 4 — Audit trail. ", True),
            ("Every state transition — owner approved, listing approved, booking confirmed, rental started, rental completed, payment recorded, cancellation with reason, reject with reason — writes an immutable row to an Activity Log with the actor’s email, a machine-readable action code, and a human-readable description. For a platform operator preparing for a compliance review, a dispute, or a stakeholder audit, this log is a complete chronicle of who did what and when.", False),
        ],
    )
    add_body(
        doc,
        [
            ("Additional safeguards: ", False),
            ("structured cancellation taxonomy", True),
            (" (customer_no_show, documents_not_verified, vehicle_unavailable, duplicate_booking, other) makes dispute resolution reviewable in aggregate rather than anecdotal; ", False),
            ("tab-mismatch login protection", True),
            (" prevents a customer account from being accidentally authenticated via the Host portal (and vice-versa); ", False),
            ("inclusive-day billing", True),
            (" ensures the host is compensated for every calendar day the car is unavailable to them, even partial handoff days.", False),
        ],
    )

    add_hr(doc)

    # ---- 5. Market Context ----
    add_heading(doc, "5. Market Context — Iloilo City", level=1)
    add_body(
        doc,
        [
            ("A candid framing first: ", True),
            ("the product’s copy is currently ", False),
            ("Philippines-general", True),
            (", not Iloilo-specific. The locale is set to en-PH, currency to PHP, the contact number uses the +63 prefix, and the testimonials reference “Metro Manila.” If this digest is paired with an Iloilo-focused pitch, the landing page and testimonial module would need a light copy pass to localize. That is a communications adjustment, not a product one — the platform itself is already city-agnostic and works wherever hosts and renters sign up.", False),
        ],
    )
    add_body(
        doc,
        [
            ("With that noted, the pain points the product addresses are ", False),
            ("strongly resonant in the Iloilo City market", True),
            (" and map directly to features already shipped:", False),
        ],
    )

    add_bullet(
        doc,
        [
            ("The trust gap in informal P2P rentals. ", True),
            ("Today, most provincial car-sharing in Iloilo happens through Facebook groups, Viber chats, and personal referrals — with no document verification, no standardized contract, and no recourse when a car comes back damaged. carBNB’s double-gated verification (host ID + license + OR/CR before approval, listing review before publication) replaces trust-by-introduction with trust-by-platform.", False),
        ],
    )
    add_bullet(
        doc,
        [
            ("Manual overhead for hosts. ", True),
            ("Iloilo’s best rental hosts today run their inventory in Google Sheets or scribbled notebooks — double-bookings are a constant operational risk, and reconciling who-paid-what at the end of the week eats weekend time. carBNB collapses all of that into one dashboard: live availability, automatic conflict prevention, a running earnings total, and a pre-filled cancellation reason dropdown for when things go sideways.", False),
        ],
    )
    add_bullet(
        doc,
        [
            ("Opaque fees in ad-hoc arrangements. ", True),
            ("Informal rentals leak margin — sometimes the “fee” is a percentage, sometimes a flat peso figure, sometimes a favor. carBNB shows the same three-line split (customer total → platform fee → owner payout) on both the renter’s dialog and the host’s booking detail. No side is ever guessing.", False),
        ],
    )
    add_bullet(
        doc,
        [
            ("Discovery friction for weekend trips. ", True),
            ("Iloilo’s rental demand is heavily weekend-weighted: road trips to Guimaras, Antique, or Boracay-via-Caticlan, family out-of-town for fiestas. A public browse page with location filtering and a 90-day availability map replaces “does your cousin know anyone with a Fortuner available next weekend?”", False),
        ],
    )
    add_bullet(
        doc,
        [
            ("Cash-economy payment reality. ", True),
            ("Gateway-based payment doesn’t fit the market yet; most handoffs are still cash on pickup. The platform’s ", False),
            ("Mark as Paid", True),
            (" flow treats cash as first-class — it’s not a workaround but a recorded event with method, timestamp, admin actor, and optional receipt notes. When payment gateways become culturally appropriate, the system extends; until then, it fits today’s behavior honestly.", False),
        ],
    )
    add_bullet(
        doc,
        [
            ("Accountability in disputes. ", True),
            ("In an informal rental scene, “the car came back with a scratch” becomes a he-said/she-said. The activity log + structured cancellation reasons + locked-in booking totals give hosts, renters, and the platform operator a shared, timestamped record when disagreements escalate.", False),
        ],
    )
    add_bullet(
        doc,
        [
            ("Idle-vehicle economics. ", True),
            ("Many Iloilo households — especially OFW families and middle-class professionals — own a second car that sits in the garage five days out of seven. carBNB’s hero copy speaks directly to this: ", False),
            ("“Rent a car or turn your idle vehicle into curated income.”", True),
            (" The host dashboard’s earnings tile makes that promise measurable.", False),
        ],
    )

    add_hr(doc)
    add_body(
        doc,
        [
            ("End of digest. For a technical accompaniment — architecture, data model, deployment pipeline — see BACKLOG.md and the project memory notes.", False),
        ],
    )

    out_path = Path(__file__).resolve().parent.parent / "CarBNB-Product-Digest.docx"
    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    build()
